import time

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import mplfinance as mpf
import pandas as pd
import numpy as np


# author 刘钰
class Report:

    # 生成报告
    @staticmethod
    def report(array, year, month, day, uid):
        document = Document()  # 实例化Document
        document.styles['Normal'].font.name = u'微软雅黑'
        document.styles['Normal'].font.size = Pt(9)
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        # 添加页眉
        header = document.sections[0].header
        p = header.paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head = str(year)+"年"+str(month)+"月"+str(day)+"日技术分析判断报告"
        font = p.add_run(head).font
        font.size = Pt(22)
        font.bold = True

        # 绘制k线图
        df = pd.DataFrame(pd.read_excel("data/testData.xlsx"))
        df.columns = ['datetime', 'Open', 'High', 'Low', 'Close']
        time1 = pd.Timestamp(year, month, day)
        df = df.loc[df['datetime'] <= time1]
        df = df.set_index('datetime')
        mpf.plot(
            data=df,
            type="candle",
            title="",
            ylabel="",
            style="binance",
            mav=(5, 10),
            volume=False,
            savefig="result/mplfinance.png"
        )

        # 添加图片
        document.add_picture("result/mplfinance.png", width=Inches(4))
        # 图片居中
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 创建表格
        table = document.add_table(rows=9, cols=0, style='Table Grid')
        table.add_column(width=Cm(4.5))
        table.add_column(width=Cm(18.5))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 填充表格内容
        header = np.asarray(["品种名称", "趋势", "N日均线价格穿越情况", "M日、N日均线交叉和穿越情况", "RSI指标", "RSI指标背离情况", "MACD穿越情况", "形态提示", "前期低点、高点是否突破"])
        for i in range(0, len(header)):
            table.cell(i, 0).text = header[i]

        # text
        table.cell(0, 1).text = "USDCNY"  # 品种名称
        table.cell(1, 1).text = array[0][0]  # 趋势
        # N日均线价格穿越情况
        string = ""
        for i in range(0, len(array[1])):
            if i == (len(array[1]) - 1):
                string += array[1][i]
            else:
                string += array[1][i] + "\n"
        table.cell(2, 1).text = string
        # M日、N日均线交叉和穿越情况
        string = ""
        for i in range(0, len(array[2])):
            if i == (len(array[2]) - 1):
                string += array[2][i]
            else:
                string += array[2][i] + "\n"
        table.cell(3, 1).text = string

        # RSI指标
        string = ""
        for i in range(0, len(array[3]), 2):
            if i != len(array[3]) - 1:
                string += array[3][i] + "       " + array[3][i + 1] + "\n"
            else:
                string += array[3][i]
        table.cell(4, 1).text = string
        # RSI指标背离情况
        string = ""
        for i in range(0, len(array[4])):
            if i == (len(array[4]) - 1):
                string += array[4][i]
            else:
                string += array[4][i] + "\n"
        table.cell(5, 1).text = string
        # MACD穿越情况
        string = ""
        for i in range(0, len(array[5]), 2):
            if i != len(array[5]) - 1:
                string += array[5][i] + "       " + array[5][i + 1] + "\n"
            else:
                string += array[5][i]
        table.cell(6, 1).text = string
        # 形态提示
        string = ""
        for i in range(0, len(array[6]), 3):
            if len(array[6]) - 1 - i >= 3:
                string += array[6][i] + "       " + array[6][i + 1] + "       " + array[6][i + 2] + "\n"
            elif len(array[6]) - 1 - i == 2:
                string += array[6][i] + "       " + array[6][i + 1] + "       " + array[6][i + 2]
            elif len(array[6]) - 1 - i == 1:
                string += array[6][i] + "       " + array[6][i + 1]
            else :
                string += array[6][i]
        table.cell(7, 1).text = string
        # 前期低点、高点是否突破
        string = array[7][0] + "            " + array[7][1]
        # print(string)
        table.cell(8, 1).text = string
        # 保存docx
        # document.save("result/report_"+time.strftime("%Y_%m_%d_%H_%M_%S", time.localtime())+".docx")
        document.save("result/report_" + uid+".docx")
